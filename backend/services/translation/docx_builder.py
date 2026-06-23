import io
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _pt_from_style(style_str: str, default: float = 9.0) -> float:
    m = re.search(r"font-size\s*:\s*([\d.]+)pt", style_str or "")
    return float(m.group(1)) if m else default


def _is_bold(style_str: str, tag: str = "") -> bool:
    if tag in ("b", "strong"):
        return True
    return "font-weight:bold" in (style_str or "").replace(" ", "")


def _align_from_style(style_str: str) -> WD_ALIGN_PARAGRAPH:
    s = (style_str or "").replace(" ", "")
    if "text-align:center" in s:
        return WD_ALIGN_PARAGRAPH.CENTER
    if "text-align:right" in s:
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def _table_has_borders(table_el) -> bool:
    """Return True if this table is a data table with visible borders."""
    # Check first td/th for border:none
    first_cell = table_el.find(["td", "th"])
    if not first_cell:
        return False
    style = first_cell.get("style", "").replace(" ", "")
    if "border:none" in style or "border:0" in style:
        return False
    table_style = table_el.get("style", "").replace(" ", "")
    if "border:none" in table_style:
        return False
    # Has explicit border → data table
    return "border:" in style or "border:" in table_style


def _clear_table_borders(tbl):
    """Remove all borders from a python-docx table (table-level + every cell)."""
    # Table-level borders
    tblPr = tbl._tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

    # Cell-level borders (overrides inherited style)
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{name}")
                el.set(qn("w:val"), "none")
                el.set(qn("w:sz"), "0")
                el.set(qn("w:color"), "auto")
                tcBorders.append(el)
            existing = tcPr.find(qn("w:tcBorders"))
            if existing is not None:
                tcPr.remove(existing)
            tcPr.append(tcBorders)


def _cell_inline_text(cell, default_bold=False, default_size=10.0):
    """Walk a table cell and return list of (text, bold, italic, size) runs."""
    from bs4 import NavigableString
    runs = []

    def _walk(node, bold=default_bold, size=default_size, italic=False):
        if isinstance(node, NavigableString):
            # Normalize whitespace: collapse newlines/tabs to single space
            t = re.sub(r'\s+', ' ', str(node))
            if t.strip():
                runs.append((t, bold, italic, size))
            elif t == ' ' and runs:  # preserve single space between words
                runs.append((' ', bold, italic, size))
            return
        n_tag = getattr(node, "name", None)
        if not n_tag or n_tag in ("script", "style"):
            return
        if n_tag == "br":
            runs.append(("\n", bold, italic, size))
            return
        n_style = node.get("style", "")
        n_bold = bold or _is_bold(n_style, n_tag)
        n_size = _pt_from_style(n_style, size)
        n_italic = italic or n_tag in ("i", "em")
        for child in node.children:
            _walk(child, n_bold, n_size, n_italic)

    for child in cell.children:
        _walk(child)
    return runs


def _add_bordered_table(doc: Document, table_el):
    """Build a bordered DOCX table from a BS4 <table> element."""
    from bs4 import NavigableString
    rows = table_el.find_all("tr", recursive=True)
    if not rows:
        return
    def _row_cols(tr):
        return sum(int(c.get("colspan", 1)) for c in tr.find_all(["td", "th"]))
    num_cols = max((_row_cols(r) for r in rows), default=0)
    if num_cols == 0:
        return

    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    tbl.style = "Table Grid"

    for r_idx, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"])
        c_idx = 0
        for cell in cells:
            if c_idx >= num_cols:
                break
            style = cell.get("style", "")
            cell_bold = cell.name == "th" or "font-weight:bold" in style.replace(" ", "")
            cell_size = _pt_from_style(style, 10.0)
            cell_align = _align_from_style(style)
            try:
                dcell = tbl.cell(r_idx, c_idx)
                p = dcell.paragraphs[0]
                p.alignment = cell_align
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.5
                for text, bold, italic, size in _cell_inline_text(cell, cell_bold, cell_size):
                    run = p.add_run(text)
                    run.bold = bold
                    run.italic = italic
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(size)
            except IndexError:
                pass
            c_idx += int(cell.get("colspan", 1))


def _is_signature_table(table_el) -> bool:
    """Detect signature/stamp block: ≤4 rows, cells mostly center-aligned."""
    rows = table_el.find_all("tr", recursive=True)
    if len(rows) > 4:
        return False
    cells = table_el.find_all(["td", "th"])
    if not cells:
        return False
    center = sum(1 for c in cells if "text-align:center" in c.get("style","").replace(" ",""))
    return center >= len(cells) / 2


def _add_html_page_to_doc(doc: Document, html: str):
    """
    Convert translated HTML to DOCX.
    - Bordered tables (data tables with visible lines) → DOCX Table Grid
    - Borderless tables (layout/fields/signatures) → plain text paragraphs
    - <p>, <h1-h4> → paragraphs preserving bold/italic/size/alignment
    """
    from bs4 import BeautifulSoup, NavigableString

    SKIP_TAGS = {"style", "script", "head"}
    DEFAULT_SIZE = 9.0

    def _el_text(el) -> str:
        parts = []
        for node in el.descendants:
            if isinstance(node, NavigableString):
                t = str(node).strip()
                if t:
                    parts.append(t)
        return " ".join(parts)

    def _add_para_inline(el):
        """Add a <p>/<h*> paragraph preserving inline bold/italic spans."""
        style = el.get("style", "")
        tag = el.name
        base_bold = _is_bold(style, tag) or tag in ("h1", "h2", "h3", "h4")
        base_size = _pt_from_style(style, DEFAULT_SIZE)
        align_str = "center" if "text-align:center" in style.replace(" ", "") else \
                    "right"  if "text-align:right"  in style.replace(" ", "") else "left"

        runs = _cell_inline_text(el, default_bold=base_bold, default_size=base_size)
        if not any(t.strip() for t, *_ in runs):
            return
        p = doc.add_paragraph()
        p.alignment = _align_from_style(f"text-align:{align_str}")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        for text, bold, italic, size in runs:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)

    def _process_borderless_table(table_el):
        """Borderless layout table → paragraphs per row, preserving <br> as line breaks."""
        tbody = table_el.find("tbody") or table_el
        for tr in tbody.find_all("tr", recursive=False):
            cells = tr.find_all(["td", "th"], recursive=False)
            if not cells:
                continue

            # Detect alignment & bold from first non-empty cell
            bold = any("font-weight:bold" in c.get("style","").replace(" ","") or c.name == "th"
                       for c in cells)
            size = _pt_from_style(tr.get("style", ""), DEFAULT_SIZE)
            center_count = sum(1 for c in cells if "text-align:center" in c.get("style","").replace(" ",""))
            align_str = "center" if center_count > len(cells) / 2 else "left"

            # Build runs for entire row: cells joined, <br> preserved as newline
            all_runs = []
            for i, cell in enumerate(cells):
                if i > 0:
                    cell_text = _el_text(cell)
                    if cell_text:
                        all_runs.append(("     ", bold, False, size))
                cstyle = cell.get("style", "")
                c_bold = bold or "font-weight:bold" in cstyle.replace(" ", "") or cell.name == "th"
                c_size = _pt_from_style(cstyle, size)
                all_runs.extend(_cell_inline_text(cell, c_bold, c_size))

            if not any(t.strip() for t, *_ in all_runs):
                continue

            p = doc.add_paragraph()
            p.alignment = _align_from_style(f"text-align:{align_str}")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            for text, b, italic, sz in all_runs:
                run = p.add_run(text)
                run.bold = b
                run.italic = italic
                run.font.name = "Times New Roman"
                run.font.size = Pt(sz)

    def _process(el):
        if isinstance(el, NavigableString):
            return
        tag = el.name
        if not tag or tag in SKIP_TAGS:
            return

        if tag == "table":
            if _table_has_borders(el):
                _add_bordered_table(doc, el)
            elif _is_signature_table(el):
                _add_bordered_table(doc, el)   # creates table, then clears borders below
                _clear_table_borders(doc.tables[-1])
            else:
                _process_borderless_table(el)
            return

        if tag in ("p", "h1", "h2", "h3", "h4"):
            _add_para_inline(el)
            return

        # div and other containers → recurse
        for child in el.children:
            _process(child)

    soup = BeautifulSoup(html, "html.parser")
    for child in soup.children:
        _process(child)



def build_docx_from_translation(translation_result: dict) -> bytes:
    """
    Generate a beautifully formatted DOCX from structured translation blocks.
    Enforces consular standards: Times New Roman, 13pt body text, 14pt bold headings,
    preserves tables and alignment formatting.
    """
    is_html_mode = translation_result.get("mode") == "html"

    doc = Document()

    # A4, consular margins: Top/Bottom 1in, Left/Right 1.25in
    for section in doc.sections:
        section.page_width    = Inches(8.27)
        section.page_height   = Inches(11.69)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Default style: Times New Roman 12pt, line spacing 1.5
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(9)
    normal.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    normal.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    p_format = normal.paragraph_format
    p_format.line_spacing = 1.5
    p_format.space_after = Pt(4)
    p_format.space_before = Pt(0)

    # Process page by page
    for p_idx, page in enumerate(translation_result.get("pages", [])):
        if p_idx > 0:
            doc.add_page_break()

        if is_html_mode:
            _add_html_page_to_doc(doc, page.get("translated_html", ""))
            continue

        for block in page.get("blocks", []):
            block_type = block.get("type", "paragraph")
            
            if block_type == "table":
                cells = block.get("translated_cells", [])
                if not cells:
                    continue
                rows_count = len(cells)
                cols_count = len(cells[0]) if rows_count > 0 else 0
                if cols_count == 0:
                    continue
                
                table = doc.add_table(rows=rows_count, cols=cols_count)
                if not block.get("borderless", False):
                    table.style = 'Table Grid'
                
                for r_idx, row_data in enumerate(cells):
                    for c_idx, cell_value in enumerate(row_data):
                        val = cell_value or ""
                        cell = table.cell(r_idx, c_idx)
                        cell.text = str(val)
                        
                        # Set cell fonts
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(13)
            
            else: # paragraph
                text = block.get("translated", "")
                if not text.strip():
                    continue
                
                is_heading = block.get("is_heading", False)
                align_str = block.get("align", "left")
                is_bold = block.get("is_bold", False)
                
                p = doc.add_paragraph()
                
                # Apply alignments
                if align_str == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align_str == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # Preserve left indent from bbox x0 if present
                    bbox = block.get("bbox")
                    if bbox and len(bbox) >= 4:
                        x0 = bbox[0]
                        if x0 > 90:
                            indent_inches = (x0 - 72) / 72.0
                            p.paragraph_format.left_indent = Inches(min(3.5, max(0.0, indent_inches)))
                    
                # Split paragraph by original newlines to preserve spacing
                lines = text.split("\n")
                for l_idx, line in enumerate(lines):
                    if l_idx > 0:
                        p.add_run("\n")
                    run = p.add_run(line)
                    run.font.name = 'Times New Roman'
                    
                    if is_heading:
                        run.font.size = Pt(14)
                        run.bold = True
                    else:
                        run.font.size = Pt(13)
                        if is_bold:
                            run.bold = True
                            
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()
