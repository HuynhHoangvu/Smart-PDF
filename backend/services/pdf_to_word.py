import io
import re
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def is_inside_any_table(bbox: list, table_bboxes: list) -> bool:
    cx = (bbox[0] + bbox[2]) / 2
    cy = (bbox[1] + bbox[3]) / 2
    for tx0, ty0, tx1, ty1 in table_bboxes:
        if tx0 - 3 <= cx <= tx1 + 3 and ty0 - 3 <= cy <= ty1 + 3:
            return True
    return False

def convert_pdf_bytes_to_docx_bytes(pdf_bytes: bytes) -> bytes:
    """
    High-fidelity PDF to Word (DOCX) converter:
    - Automatically detects and reconstructs tables.
    - Uses OCR (RapidOCR) fallback for scanned image-only PDFs.
    - Preserves text fonts, alignments (left/center/right), and bold flags.
    - Formats output using Times New Roman (user standard).
    """
    blocks_data = parse_pdf_to_blocks(pdf_bytes)
    return build_docx_from_blocks(blocks_data)

def parse_pdf_to_blocks(pdf_bytes: bytes) -> dict:
    """
    Parses PDF bytes and extracts text paragraphs and tables as structured blocks per page.
    """
    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages_data = []

    ocr_engine = None

    for page_index in range(len(pdf_doc)):
        page = pdf_doc.load_page(page_index)
        page_width = page.rect.width
        
        # 1. Table Detection
        tables = page.find_tables()
        table_bboxes = [list(t.bbox) for t in tables]
        extracted_tables = []
        for t in tables:
            extracted_tables.append({
                "type": "table",
                "bbox": list(t.bbox),
                "cells": t.extract()
            })

        # 2. Text Extraction
        raw_text = page.get_text("text").strip()
        needs_ocr = len(raw_text) < 30
        paragraphs = []

        if needs_ocr:
            if ocr_engine is None:
                from rapidocr_onnxruntime import RapidOCR
                ocr_engine = RapidOCR()
            
            # Render page at 2x zoom for crisp OCR
            zoom = 2.0
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("png")
            
            ocr_result, _ = ocr_engine(img_bytes)
            if ocr_result:
                for box, text, conf in ocr_result:
                    t_str = text.strip()
                    if not t_str:
                        continue
                    xs = [pt[0] for pt in box]
                    ys = [pt[1] for pt in box]
                    x0 = min(xs) / zoom
                    y0 = min(ys) / zoom
                    x1 = max(xs) / zoom
                    y1 = max(ys) / zoom
                    
                    if is_inside_any_table([x0, y0, x1, y1], table_bboxes):
                        continue

                    # Heuristics
                    h_val = y1 - y0
                    font_size_pt = round(h_val * 0.75, 2)
                    is_bold = font_size_pt > 14
                    is_heading = font_size_pt >= 14

                    block_width = x1 - x0
                    page_mid = page_width / 2

                    align = "left"
                    margin_left = x0
                    margin_right = page_width - x1
                    if is_heading and abs(margin_left - margin_right) < (page_width * 0.1) and block_width < (page_width * 0.8):
                        align = "center"
                    elif x0 > page_mid * 1.05 and block_width < (page_width * 0.5):
                        align = "right"
                    elif abs(margin_left - margin_right) < (page_width * 0.05) and block_width < (page_width * 0.6):
                        align = "center"

                    paragraphs.append({
                        "type": "paragraph",
                        "text": t_str,
                        "bbox": [x0, y0, x1, y1],
                        "align": align,
                        "is_bold": is_bold,
                        "is_heading": is_heading,
                        "font_size": font_size_pt
                    })
        else:
            # Native text extraction
            raw = page.get_text("dict")
            all_sizes = []
            for block in raw["blocks"]:
                if "lines" not in block:
                    continue
                for line in block["lines"]:
                    for span in line["spans"]:
                        if span["text"].strip():
                            all_sizes.append(span["size"])

            avg_size = sum(all_sizes) / len(all_sizes) if all_sizes else 11.0
            heading_threshold = avg_size * 1.15

            for block in raw["blocks"]:
                if "lines" not in block:
                    continue
                bbox = list(block["bbox"])
                if is_inside_any_table(bbox, table_bboxes):
                    continue

                lines_text = []
                max_font_size = 0.0
                block_is_bold = False

                for line in block["lines"]:
                    line_parts = []
                    for span in line["spans"]:
                        t = span["text"]
                        if not t:
                            continue
                        line_parts.append(t)
                        if span["size"] > max_font_size:
                            max_font_size = span["size"]
                        if span["flags"] & 16:
                            block_is_bold = True
                    if line_parts:
                        lines_text.append(" ".join(line_parts))

                block_text = "\n".join(lines_text).strip()
                if not block_text:
                    continue

                is_heading = (max_font_size >= heading_threshold) or block_is_bold

                x0, y0, x1, y1 = bbox
                block_width = x1 - x0
                page_mid = page_width / 2

                align = "left"
                margin_left = x0
                margin_right = page_width - x1
                if is_heading and abs(margin_left - margin_right) < (page_width * 0.15) and block_width < (page_width * 0.8):
                    align = "center"
                elif x0 > page_mid * 1.05 and block_width < (page_width * 0.5):
                    align = "right"
                elif abs(margin_left - margin_right) < (page_width * 0.05) and block_width < (page_width * 0.6):
                    align = "center"

                paragraphs.append({
                    "type": "paragraph",
                    "text": block_text,
                    "bbox": bbox,
                    "align": align,
                    "is_bold": block_is_bold,
                    "is_heading": is_heading,
                    "font_size": max_font_size
                })

        # Merge and sort page content (paragraphs and tables) visually top-to-bottom
        all_blocks = paragraphs + extracted_tables
        all_blocks.sort(key=lambda b: b["bbox"][1])
        pages_data.append({
            "page_num": page_index + 1,
            "blocks": all_blocks
        })

    pdf_doc.close()
    return {
        "pages": pages_data
    }

def build_docx_from_blocks(blocks_data: dict) -> bytes:
    """
    Builds a beautifully formatted DOCX from structured blocks data.
    """
    word_doc = Document()

    # Margins setup
    for section in word_doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configure Normal style
    style = word_doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13) # Enforce Times New Roman 13pt body, 14pt bold headings

    p_format = style.paragraph_format
    p_format.line_spacing = 1.25
    p_format.space_after = Pt(4)
    p_format.space_before = Pt(0)

    for p_idx, page in enumerate(blocks_data.get("pages", [])):
        if p_idx > 0:
            word_doc.add_page_break()

        for block in page.get("blocks", []):
            if block["type"] == "table":
                cells = block.get("cells", [])
                if not cells:
                    continue
                rows_count = len(cells)
                cols_count = len(cells[0]) if rows_count > 0 else 0
                if cols_count == 0:
                    continue
                
                table = word_doc.add_table(rows=rows_count, cols=cols_count)
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(cells):
                    for c_idx, cell_value in enumerate(row_data):
                        cell_text = cell_value or ""
                        cell = table.cell(r_idx, c_idx)
                        cell.text = str(cell_text)
                        
                        # Apply font styles
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.name = 'Times New Roman'
                                r.font.size = Pt(13)
            else:
                p = word_doc.add_paragraph()
                align_str = block.get("align", "left")
                if align_str == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align_str == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # Preserve indentations based on x0 (PDF coordinate points)
                    bbox = block.get("bbox")
                    if bbox and len(bbox) >= 4:
                        x0 = bbox[0]
                        # PDF points are 72 per inch. Default page left margin is 1.0 inch (72 points).
                        # Apply relative indent if it's indented further to the right.
                        if x0 > 90:
                            indent_inches = (x0 - 72) / 72.0
                            p.paragraph_format.left_indent = Inches(min(3.5, max(0.0, indent_inches)))

                lines = block.get("text", "").split("\n")
                for l_idx, line in enumerate(lines):
                    if l_idx > 0:
                        p.add_run("\n")
                    run = p.add_run(line)
                    run.font.name = 'Times New Roman'
                    if block.get("is_heading"):
                        run.font.size = Pt(14)
                        run.bold = True
                    else:
                        run.font.size = Pt(13)
                        if block.get("is_bold"):
                            run.bold = True

    output = io.BytesIO()
    word_doc.save(output)
    output.seek(0)
    return output.read()
